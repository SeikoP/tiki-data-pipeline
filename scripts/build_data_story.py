"""
Script để phân tích database và xây dựng câu chuyện dữ liệu (Data Story)
Tạo file docx chứa câu chuyện dữ liệu của dự án Tiki Data Pipeline
"""

import os
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

try:
    import psycopg2
    from psycopg2.extras import RealDictCursor
except ImportError:
    print("❌ Cần cài đặt psycopg2-binary: pip install psycopg2-binary")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("❌ Cần cài đặt python-docx: pip install python-docx")
    sys.exit(1)

# Thêm src vào path để import modules
PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT / "src"))

# Đọc thông tin từ .env
def load_env_config():
    """Đọc cấu hình từ file .env"""
    env_file = PROJECT_ROOT / ".env"
    config = {
        "host": "localhost",
        "port": 5432,
        "database": "crawl_data",
        "user": "postgres",
        "password": "postgres",
    }
    
    if env_file.exists():
        with open(env_file, encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith("#") and "=" in line:
                    key, value = line.split("=", 1)
                    key = key.strip()
                    value = value.strip().strip('"').strip("'")
                    
                    if key == "POSTGRES_HOST":
                        config["host"] = value
                    elif key == "POSTGRES_PORT":
                        config["port"] = int(value)
                    elif key == "POSTGRES_USER":
                        config["user"] = value
                    elif key == "POSTGRES_PASSWORD":
                        config["password"] = value
                    elif key == "POSTGRES_DB" and value == "crawl_data":
                        config["database"] = value
    
    # Override với environment variables nếu có
    config["host"] = os.getenv("POSTGRES_HOST", config["host"])
    config["port"] = int(os.getenv("POSTGRES_PORT", config["port"]))
    config["user"] = os.getenv("POSTGRES_USER", config["user"])
    config["password"] = os.getenv("POSTGRES_PASSWORD", config["password"])
    config["database"] = os.getenv("POSTGRES_DB", config["database"])
    
    return config


def connect_database(config: dict[str, Any]):
    """Kết nối đến PostgreSQL database"""
    try:
        conn = psycopg2.connect(
            host=config["host"],
            port=config["port"],
            database=config["database"],
            user=config["user"],
            password=config["password"],
            connect_timeout=10,
        )
        print(f"✅ Đã kết nối đến database: {config['database']}")
        return conn
    except Exception as e:
        print(f"❌ Lỗi kết nối database: {e}")
        sys.exit(1)


def analyze_database(conn) -> dict[str, Any]:
    """Phân tích dữ liệu trong database"""
    stats = {}
    
    with conn.cursor(cursor_factory=RealDictCursor) as cur:
        # 1. Tổng quan
        print("📊 Đang phân tích tổng quan...")
        cur.execute("""
            SELECT 
                (SELECT COUNT(*) FROM categories) as total_categories,
                (SELECT COUNT(*) FROM products) as total_products,
                (SELECT COUNT(DISTINCT category_url) FROM products WHERE category_url IS NOT NULL) as categories_with_products
        """)
        stats["overview"] = dict(cur.fetchone())
        
        # 2. Thống kê categories
        print("📁 Đang phân tích categories...")
        cur.execute("""
            SELECT 
                COUNT(*) as total,
                COUNT(DISTINCT level) as distinct_levels,
                MIN(level) as min_level,
                MAX(level) as max_level,
                AVG(product_count) as avg_products_per_category,
                SUM(product_count) as total_product_count_in_categories
            FROM categories
        """)
        stats["categories"] = dict(cur.fetchone())
        
        # 3. Thống kê products
        print("🛍️ Đang phân tích products...")
        cur.execute("""
            SELECT 
                COUNT(*) as total,
                COUNT(DISTINCT category_url) as categories_covered,
                COUNT(DISTINCT brand) as distinct_brands,
                COUNT(DISTINCT seller_id) as distinct_sellers,
                COUNT(CASE WHEN seller_is_official = TRUE THEN 1 END) as official_seller_products,
                AVG(sales_count) as avg_sales_count,
                MAX(sales_count) as max_sales_count,
                MIN(sales_count) as min_sales_count,
                AVG(price) as avg_price,
                MAX(price) as max_price,
                MIN(price) as min_price,
                AVG(rating_average) as avg_rating,
                AVG(review_count) as avg_reviews,
                COUNT(CASE WHEN stock_available = TRUE THEN 1 END) as in_stock_count,
                COUNT(CASE WHEN discount_percent > 0 THEN 1 END) as discounted_products,
                AVG(discount_percent) as avg_discount_percent
            FROM products
        """)
        stats["products"] = dict(cur.fetchone())
        
        # 4. Top categories theo số lượng sản phẩm
        print("🏆 Đang lấy top categories...")
        cur.execute("""
            SELECT 
                c.name,
                c.level,
                COUNT(p.id) as product_count,
                AVG(p.price) as avg_price,
                AVG(p.sales_count) as avg_sales
            FROM categories c
            LEFT JOIN products p ON c.url = p.category_url
            GROUP BY c.id, c.name, c.level
            HAVING COUNT(p.id) > 0
            ORDER BY product_count DESC
            LIMIT 10
        """)
        stats["top_categories"] = [dict(row) for row in cur.fetchall()]
        
        # 5. Top products theo sales_count
        print("⭐ Đang lấy top products...")
        cur.execute("""
            SELECT 
                product_id,
                name,
                price,
                sales_count,
                rating_average,
                review_count,
                brand,
                seller_name
            FROM products
            WHERE sales_count IS NOT NULL
            ORDER BY sales_count DESC
            LIMIT 10
        """)
        stats["top_products"] = [dict(row) for row in cur.fetchall()]
        
        # 6. Phân bố giá
        print("💰 Đang phân tích phân bố giá...")
        cur.execute("""
            SELECT 
                CASE 
                    WHEN price < 100000 THEN 'Dưới 100k'
                    WHEN price < 500000 THEN '100k - 500k'
                    WHEN price < 1000000 THEN '500k - 1M'
                    WHEN price < 5000000 THEN '1M - 5M'
                    ELSE 'Trên 5M'
                END as price_range,
                COUNT(*) as count,
                AVG(price) as avg_price
            FROM products
            WHERE price IS NOT NULL
            GROUP BY price_range
            ORDER BY MIN(price)
        """)
        stats["price_distribution"] = [dict(row) for row in cur.fetchall()]
        
        # 7. Phân bố rating
        print("⭐ Đang phân tích rating...")
        cur.execute("""
            SELECT 
                CASE 
                    WHEN rating_average IS NULL THEN 'Chưa có rating'
                    WHEN rating_average < 3.0 THEN 'Dưới 3.0'
                    WHEN rating_average < 4.0 THEN '3.0 - 4.0'
                    WHEN rating_average < 4.5 THEN '4.0 - 4.5'
                    ELSE 'Trên 4.5'
                END as rating_range,
                COUNT(*) as count
            FROM products
            GROUP BY rating_range
            ORDER BY MIN(COALESCE(rating_average, 0))
        """)
        stats["rating_distribution"] = [dict(row) for row in cur.fetchall()]
        
        # 8. Thống kê theo brand
        print("🏷️ Đang phân tích brands...")
        cur.execute("""
            SELECT 
                brand,
                COUNT(*) as product_count,
                AVG(price) as avg_price,
                AVG(sales_count) as avg_sales,
                AVG(rating_average) as avg_rating
            FROM products
            WHERE brand IS NOT NULL
            GROUP BY brand
            ORDER BY product_count DESC
            LIMIT 10
        """)
        stats["top_brands"] = [dict(row) for row in cur.fetchall()]
        
        # 9. Phân tích computed fields
        print("📈 Đang phân tích computed fields...")
        cur.execute("""
            SELECT 
                AVG(estimated_revenue) as avg_revenue,
                SUM(estimated_revenue) as total_revenue,
                AVG(popularity_score) as avg_popularity,
                MAX(popularity_score) as max_popularity,
                AVG(value_score) as avg_value_score,
                MAX(value_score) as max_value_score,
                COUNT(CASE WHEN price_category = 'budget' THEN 1 END) as budget_count,
                COUNT(CASE WHEN price_category = 'mid-range' THEN 1 END) as midrange_count,
                COUNT(CASE WHEN price_category = 'premium' THEN 1 END) as premium_count,
                COUNT(CASE WHEN price_category = 'luxury' THEN 1 END) as luxury_count
            FROM products
            WHERE estimated_revenue IS NOT NULL
        """)
        stats["computed_fields"] = dict(cur.fetchone())
        
        # 10. Mối quan hệ giữa giá và doanh số
        print("🔗 Đang phân tích mối quan hệ giá-doanh số...")
        cur.execute("""
            SELECT 
                CASE 
                    WHEN price < 100000 THEN 'Dưới 100k'
                    WHEN price < 500000 THEN '100k - 500k'
                    WHEN price < 1000000 THEN '500k - 1M'
                    WHEN price < 5000000 THEN '1M - 5M'
                    ELSE 'Trên 5M'
                END as price_range,
                AVG(sales_count) as avg_sales,
                AVG(rating_average) as avg_rating,
                COUNT(*) as product_count
            FROM products
            WHERE price IS NOT NULL AND sales_count IS NOT NULL
            GROUP BY price_range
            ORDER BY MIN(price)
        """)
        stats["price_sales_relationship"] = [dict(row) for row in cur.fetchall()]
        
        # 11. Mối quan hệ giữa discount và sales
        print("💰 Đang phân tích tác động của discount...")
        cur.execute("""
            SELECT 
                CASE 
                    WHEN discount_percent = 0 OR discount_percent IS NULL THEN 'Không giảm giá'
                    WHEN discount_percent < 10 THEN 'Giảm dưới 10%'
                    WHEN discount_percent < 20 THEN 'Giảm 10-20%'
                    WHEN discount_percent < 30 THEN 'Giảm 20-30%'
                    ELSE 'Giảm trên 30%'
                END as discount_range,
                AVG(sales_count) as avg_sales,
                COUNT(*) as product_count,
                AVG(rating_average) as avg_rating
            FROM products
            WHERE sales_count IS NOT NULL
            GROUP BY discount_range
            ORDER BY MIN(COALESCE(discount_percent, 0))
        """)
        stats["discount_impact"] = [dict(row) for row in cur.fetchall()]
        
        # 12. Top products theo computed fields
        print("⭐ Đang lấy top products theo computed fields...")
        cur.execute("""
            SELECT 
                product_id,
                name,
                price,
                sales_count,
                rating_average,
                popularity_score,
                value_score,
                estimated_revenue,
                discount_percent
            FROM products
            WHERE popularity_score IS NOT NULL
            ORDER BY popularity_score DESC
            LIMIT 5
        """)
        stats["top_by_popularity"] = [dict(row) for row in cur.fetchall()]
        
        cur.execute("""
            SELECT 
                product_id,
                name,
                price,
                sales_count,
                rating_average,
                popularity_score,
                value_score,
                estimated_revenue
            FROM products
            WHERE value_score IS NOT NULL
            ORDER BY value_score DESC
            LIMIT 5
        """)
        stats["top_by_value"] = [dict(row) for row in cur.fetchall()]
    
    return stats


def safe_format(value: Any, format_str: str = ",.0f") -> str:
    """Format số an toàn, xử lý None"""
    if value is None:
        return "N/A"
    try:
        return f"{value:{format_str}}"
    except (ValueError, TypeError):
        return str(value) if value else "N/A"


def create_document(stats: dict[str, Any], output_path: Path):
    """Tạo file docx với câu chuyện dữ liệu"""
    doc = Document()
    
    # Cấu hình font cho tiếng Việt
    def set_vietnamese_font(run):
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    
    # Title
    title = doc.add_heading("Câu Chuyện Dữ Liệu - Tiki Data Pipeline", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Subtitle
    subtitle = doc.add_paragraph(f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()  # Spacing
    
    # ============================================
    # PHẦN BỐI CẢNH VÀ GIỚI THIỆU
    # ============================================
    
    doc.add_heading("Bối Cảnh: Thị Trường Thương Mại Điện Tử Việt Nam", 1)
    
    # Bối cảnh thị trường
    context_market = doc.add_paragraph()
    context_market.add_run("Thương mại điện tử Việt Nam đang phát triển mạnh mẽ, dự kiến đạt 49 tỷ USD vào năm 2025. ")
    context_market.add_run("Tiki.vn là một trong những nền tảng thương mại điện tử hàng đầu, được thành lập từ năm 2010 với hàng triệu sản phẩm đa dạng. ")
    context_market.add_run("Dữ liệu từ Tiki.vn phản ánh xu hướng mua sắm, hành vi tiêu dùng và cấu trúc thị trường, có giá trị nghiên cứu cao.")
    
    for run in context_market.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # Ý nghĩa của dataset
    context_meaning = doc.add_paragraph()
    context_meaning.add_run("Dataset này không chỉ là danh sách sản phẩm, mà là cửa sổ để hiểu về thị trường thương mại điện tử Việt Nam. ")
    context_meaning.add_run("Từ dataset có thể khám phá: xu hướng tiêu dùng, cấu trúc thị trường, hành vi mua sắm, sự cạnh tranh giữa các thương hiệu, và giá trị thị trường.")
    
    for run in context_meaning.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # Lý do chọn đề tài
    doc.add_heading("Lý Do Chọn Đề Tài", 1)
    
    reason_intro = doc.add_paragraph()
    reason_intro.add_run("Việc xây dựng dataset và phân tích dữ liệu từ Tiki.vn được lựa chọn dựa trên những lý do sau:")
    
    for run in reason_intro.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(12)
    
    reasons = [
        ("Tầm quan trọng của thương mại điện tử", 
         "Thương mại điện tử đang là xu hướng tất yếu của thời đại số. Việc hiểu rõ thị trường này không chỉ quan trọng đối với các doanh nghiệp, mà còn có giá trị nghiên cứu và học thuật cao. Dữ liệu từ Tiki.vn cung cấp cái nhìn toàn diện về hành vi tiêu dùng, xu hướng thị trường và cấu trúc kinh doanh."),
        
        ("Giá trị thực tiễn của dữ liệu", 
         "Khác với các dataset mẫu trong sách giáo khoa, dữ liệu từ Tiki.vn là dữ liệu thực tế từ thị trường. Dataset này có thể được sử dụng để nghiên cứu thị trường, phân tích cạnh tranh, dự đoán xu hướng. Đây là cơ hội để làm việc với dữ liệu thực tế và áp dụng kiến thức đã học."),
        
        ("Thách thức kỹ thuật và cơ hội học hỏi", 
         "Xây dựng hệ thống thu thập dữ liệu từ website động như Tiki.vn là một thách thức kỹ thuật. Dự án này đòi hỏi kiến thức về web scraping, xử lý dữ liệu, thiết kế database, và phân tích dữ liệu. Đây là cơ hội để áp dụng và nâng cao các kỹ năng trong lĩnh vực Data Engineering và Data Analytics."),
        
        ("Tính mới và đóng góp", 
         "Mặc dù có nhiều nghiên cứu về thương mại điện tử, nhưng việc xây dựng một dataset toàn diện và có hệ thống từ Tiki.vn vẫn còn hạn chế. Dataset này có thể đóng góp cho cộng đồng nghiên cứu, các nhà phân tích dữ liệu, và những người quan tâm đến thị trường thương mại điện tử Việt Nam."),
        
        ("Ứng dụng trong giáo dục và nghiên cứu", 
         "Dataset này có thể được sử dụng như một case study trong giáo dục về Data Engineering, Data Analytics, và Business Intelligence. Nó cung cấp một ví dụ thực tế về cách thu thập, xử lý, và phân tích dữ liệu từ nguồn thực tế, giúp sinh viên và nhà nghiên cứu hiểu rõ hơn về quy trình làm việc với dữ liệu thực tế."),
        
        ("Tiềm năng mở rộng", 
         "Dự án này có tiềm năng mở rộng lớn. Có thể phát triển thành một hệ thống monitoring thị trường, một công cụ phân tích cạnh tranh, hoặc một nền tảng cung cấp dữ liệu cho các ứng dụng khác. Dataset có thể được cập nhật định kỳ để theo dõi sự thay đổi của thị trường.")
    ]
    
    for idx, (title, content) in enumerate(reasons, 1):
        reason_heading = doc.add_heading(f"{idx}. {title}", 2)
        for run in reason_heading.runs:
            set_vietnamese_font(run)
        
        reason_para = doc.add_paragraph(content)
        for run in reason_para.runs:
            set_vietnamese_font(run)
            run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # Lời mở đầu
    doc.add_heading("Lời Mở Đầu: Câu Chuyện Từ Dữ Liệu", 1)
    intro = doc.add_paragraph()
    intro.add_run("Đằng sau mỗi con số là một câu chuyện. ")
    intro.add_run("Đằng sau mỗi sản phẩm là một lựa chọn của người tiêu dùng. ")
    intro.add_run("Đằng sau mỗi danh mục là một xu hướng thị trường. ")
    intro.add_run("Tài liệu này trình bày những câu chuyện được khám phá từ dữ liệu thu thập từ Tiki.vn, một trong những nền tảng thương mại điện tử hàng đầu Việt Nam.")
    
    for run in intro.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # Giới thiệu về dataset
    doc.add_heading("Về Dataset", 2)
    dataset_intro = doc.add_paragraph()
    dataset_intro.add_run("Dataset này chứa thông tin về hàng nghìn sản phẩm từ Tiki.vn, được thu thập và xử lý một cách có hệ thống. ")
    dataset_intro.add_run("Mỗi sản phẩm trong dataset bao gồm thông tin chi tiết về: tên sản phẩm, giá cả, mô tả, đánh giá của người dùng, thông tin người bán, thương hiệu, số lượng đã bán, và nhiều chỉ số phân tích khác. ")
    dataset_intro.add_run("Mỗi dòng dữ liệu không chỉ phản ánh thông tin về sản phẩm, mà còn cho thấy về thị trường, về hành vi mua sắm của người tiêu dùng, và về những xu hướng đang diễn ra.")
    
    for run in dataset_intro.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # Câu chuyện từ dữ liệu
    doc.add_heading("Những Câu Hỏi Nghiên Cứu", 2)
    story_intro = doc.add_paragraph()
    story_intro.add_run("Khi bắt đầu với dataset này, có nhiều câu hỏi nghiên cứu được đặt ra. ")
    story_intro.add_run("Dữ liệu sẽ giúp trả lời những câu hỏi đó. ")
    story_intro.add_run("Dưới đây là những vấn đề có thể khám phá từ dataset:")
    
    for run in story_intro.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(12)
    
    story_points = [
        "Thị trường Tiki có quy mô như thế nào? Có bao nhiêu sản phẩm và danh mục?",
        "Người tiêu dùng đang mua gì? Sản phẩm nào được mua nhiều nhất?",
        "Giá cả trên thị trường phân bố như thế nào? Người tiêu dùng thường mua ở mức giá nào?",
        "Thương hiệu nào đang dẫn đầu? Ai là người bán tốt nhất?",
        "Người tiêu dùng đánh giá sản phẩm như thế nào? Điểm số và review phản ánh điều gì?",
        "Giảm giá có thực sự ảnh hưởng đến doanh số không?"
    ]
    
    for point in story_points:
        p = doc.add_paragraph(point, style="List Bullet")
        for run in p.runs:
            set_vietnamese_font(run)
            run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # 1. Tổng quan
    doc.add_heading("1. Tổng Quan Dữ Liệu: Mẫu Nghiên Cứu", 1)
    overview = stats["overview"]
    overview_text = doc.add_paragraph()
    overview_text.add_run("Sau quá trình thu thập và xử lý, dataset đã được xây dựng với quy mô đáng kể. ")
    overview_text.add_run(f"Dataset hiện tại bao gồm ")
    overview_text.add_run(f"{overview['total_products']:,}").bold = True
    overview_text.add_run(" sản phẩm từ ")
    overview_text.add_run(f"{overview['total_categories']:,}").bold = True
    overview_text.add_run(" danh mục khác nhau. ")
    overview_text.add_run(f"Trong đó, có ")
    overview_text.add_run(f"{overview['categories_with_products']:,}").bold = True
    overview_text.add_run(" danh mục thực sự có sản phẩm. ")
    overview_text.add_run("Cần lưu ý rằng đây là một mẫu dữ liệu được thu thập, không phải toàn bộ sản phẩm trên Tiki.vn. ")
    overview_text.add_run("Tuy nhiên, với quy mô này, dataset vẫn đủ lớn và đại diện để có thể nghiên cứu về các xu hướng và đặc điểm của thị trường thương mại điện tử Việt Nam.")
    
    for run in overview_text.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(12)
    
    # Insights từ tổng quan
    insight_para = doc.add_paragraph()
    if overview["categories_with_products"] and overview["total_categories"]:
        coverage_ratio = (overview["categories_with_products"] / overview["total_categories"]) * 100
        insight_para.add_run(f"Một phát hiện thú vị: trong số ")
        insight_para.add_run(f"{overview['total_categories']:,}").bold = True
        insight_para.add_run(" danh mục được thu thập, có ")
        insight_para.add_run(f"{overview['categories_with_products']:,}").bold = True
        insight_para.add_run(f" danh mục ({coverage_ratio:.1f}%) thực sự có sản phẩm. ")
        insight_para.add_run("Điều này cho thấy một số danh mục có thể là danh mục cha (chỉ để phân loại) hoặc danh mục trống, phản ánh cách Tiki tổ chức cấu trúc sản phẩm.")
    
    for run in insight_para.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # 2. Phân tích Categories
    doc.add_heading("2. Câu Chuyện Về Danh Mục: Cấu Trúc Thị Trường", 1)
    cat_stats = stats["categories"]
    cat_text = doc.add_paragraph()
    cat_text.add_run("Danh mục sản phẩm phản ánh cách tổ chức và phân loại thị trường. ")
    cat_text.add_run("Tiki.vn tổ chức sản phẩm theo cấu trúc phân cấp đa tầng với ")
    cat_text.add_run(f"{cat_stats['distinct_levels']}").bold = True
    cat_text.add_run(" cấp độ khác nhau, từ cấp ")
    cat_text.add_run(f"{cat_stats['min_level']}").bold = True
    cat_text.add_run(" đến cấp ")
    cat_text.add_run(f"{cat_stats['max_level']}").bold = True
    cat_text.add_run(". ")
    if cat_stats["avg_products_per_category"]:
        cat_text.add_run(f"Trung bình mỗi danh mục có ")
        cat_text.add_run(f"{safe_format(cat_stats.get('avg_products_per_category'), '.0f')}").bold = True
        cat_text.add_run(" sản phẩm.")
    
    for run in cat_text.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(12)
    
    # Insights về danh mục
    if stats["top_categories"]:
        top_cat_insight = doc.add_paragraph()
        top_cat = stats["top_categories"][0] if stats["top_categories"] else None
        if top_cat:
            top_cat_insight.add_run("Danh mục có nhiều sản phẩm nhất là ")
            top_cat_insight.add_run(f'"{top_cat["name"]}"').bold = True
            top_cat_insight.add_run(f" với ")
            top_cat_insight.add_run(f"{top_cat['product_count']:,}").bold = True
            top_cat_insight.add_run(" sản phẩm. ")
            if top_cat.get("avg_price"):
                top_cat_insight.add_run(f"Giá trung bình trong danh mục này là ")
                top_cat_insight.add_run(f"{safe_format(top_cat.get('avg_price'), ',.0f')} VND").bold = True
                top_cat_insight.add_run(", cho thấy phân khúc giá của danh mục này.")
            
            for run in top_cat_insight.runs:
                set_vietnamese_font(run)
                run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # 3. Phân tích Products
    doc.add_heading("3. Câu Chuyện Về Sản Phẩm: Thị Trường Trong Lòng Bàn Tay", 1)
    prod_stats = stats["products"]
    prod_text = doc.add_paragraph()
    prod_text.add_run("Mỗi sản phẩm trong dataset phản ánh một lựa chọn của người tiêu dùng dựa trên nhu cầu, giá cả, và đánh giá. ")
    prod_text.add_run(f"Dataset bao gồm sản phẩm từ ")
    prod_text.add_run(f"{prod_stats['distinct_brands']:,}").bold = True
    prod_text.add_run(" thương hiệu khác nhau và ")
    prod_text.add_run(f"{prod_stats['distinct_sellers']:,}").bold = True
    prod_text.add_run(" người bán. ")
    prod_text.add_run("Điều này cho thấy thị trường rất đa dạng và cạnh tranh, tạo ra nhiều lựa chọn cho người tiêu dùng.")
    
    for run in prod_text.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(12)
    
    # Câu chuyện về giá cả
    if prod_stats.get("avg_price") is not None:
        price_story = doc.add_paragraph()
        price_story.add_run("Một câu hỏi nghiên cứu quan trọng: Người tiêu dùng thường mua ở mức giá nào? ")
        if stats["price_distribution"]:
            max_count_range = max(stats["price_distribution"], key=lambda x: x["count"])
            total_products = sum(r["count"] for r in stats["price_distribution"])
            max_percentage = (max_count_range["count"] / total_products) * 100 if total_products > 0 else 0
            price_story.add_run(f"Dữ liệu cho thấy ")
            price_story.add_run(f"{max_percentage:.1f}%").bold = True
            price_story.add_run(f" sản phẩm nằm trong khoảng ")
            price_story.add_run(f'"{max_count_range["price_range"]}"').bold = True
            price_story.add_run(". ")
            price_story.add_run("Điều này phản ánh phân khúc giá mà người tiêu dùng Việt Nam thường lựa chọn khi mua sắm online.")
        else:
            price_story.add_run(f"Giá trung bình là ")
            price_story.add_run(f"{safe_format(prod_stats.get('avg_price'), ',.0f')} VND").bold = True
            price_story.add_run(", cho thấy mức giá phổ biến trên thị trường.")
        
        for run in price_story.runs:
            set_vietnamese_font(run)
            run.font.size = Pt(12)
    
    # Câu chuyện về sản phẩm bán chạy
    if stats["top_products"]:
        top_prod_story = doc.add_paragraph()
        top_prod = stats["top_products"][0] if stats["top_products"] else None
        if top_prod:
            top_prod_story.add_run("Sản phẩm nào được mua nhiều nhất? ")
            top_prod_story.add_run(f'"{top_prod["name"][:60]}"').bold = True
            if top_prod.get("sales_count"):
                top_prod_story.add_run(f" đã bán được ")
                top_prod_story.add_run(f"{safe_format(top_prod.get('sales_count'), ',')}").bold = True
                top_prod_story.add_run(" sản phẩm. ")
            top_prod_story.add_run("Điều này cho thấy sản phẩm này đáp ứng được nhu cầu và sở thích của đông đảo người tiêu dùng.")
            
            for run in top_prod_story.runs:
                set_vietnamese_font(run)
                run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # 4. Phân tích Brands
    if stats["top_brands"]:
        doc.add_heading("4. Câu Chuyện Về Thương Hiệu: Ai Đang Dẫn Đầu?", 1)
        brand_text = doc.add_paragraph()
        brand_text.add_run("Thương hiệu đóng vai trò quan trọng trong quyết định mua sắm của người tiêu dùng. ")
        brand_text.add_run("Thương hiệu không chỉ là tên gọi, mà còn là lời hứa về chất lượng và giá trị. ")
        brand_text.add_run("Trên Tiki, có rất nhiều thương hiệu cạnh tranh với nhau để thu hút người mua, tạo nên một thị trường đa dạng và sôi động.")
        
        for run in brand_text.runs:
            set_vietnamese_font(run)
            run.font.size = Pt(12)
        
        # Câu chuyện về thương hiệu hàng đầu
        top_brand = stats["top_brands"][0] if stats["top_brands"] else None
        if top_brand:
            brand_story = doc.add_paragraph()
            brand_story.add_run("Thương hiệu có nhiều sản phẩm nhất là ")
            brand_story.add_run(f'"{top_brand["brand"]}"').bold = True
            brand_story.add_run(f" với ")
            brand_story.add_run(f"{top_brand['product_count']:,}").bold = True
            brand_story.add_run(" sản phẩm. ")
            brand_story.add_run("Điều này cho thấy thương hiệu này đã xây dựng được một danh mục sản phẩm đa dạng và có vị thế mạnh trên thị trường.")
            
            for run in brand_story.runs:
                set_vietnamese_font(run)
                run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # 5. Câu chuyện về giá trị thị trường
    if stats.get("computed_fields"):
        doc.add_heading("5. Câu Chuyện Về Giá Trị Thị Trường", 1)
        computed = stats["computed_fields"]
        
        # Estimated Revenue - câu chuyện về quy mô
        if computed.get("total_revenue"):
            revenue_story = doc.add_paragraph()
            revenue_story.add_run("Một câu hỏi nghiên cứu quan trọng: Thị trường này có giá trị bao nhiêu? ")
            revenue_story.add_run("Từ dataset, tổng doanh thu ước tính là ")
            revenue_story.add_run(f"{safe_format(computed.get('total_revenue') / 1000000000, '.2f')} tỷ VND").bold = True
            revenue_story.add_run(". ")
            revenue_story.add_run("Con số này phản ánh quy mô và tiềm năng của thị trường thương mại điện tử Việt Nam.")
            
            for run in revenue_story.runs:
                set_vietnamese_font(run)
                run.font.size = Pt(12)
        
        doc.add_paragraph()  # Spacing
    
    # 6. Câu chuyện về mối quan hệ giá và doanh số
    if stats.get("price_sales_relationship"):
        doc.add_heading("6. Câu Chuyện: Giá Nào Bán Chạy Nhất?", 1)
        relationship_story = doc.add_paragraph()
        relationship_story.add_run("Một câu hỏi nghiên cứu quan trọng: Ở mức giá nào thì sản phẩm bán chạy nhất? ")
        relationship_story.add_run("Đây là câu hỏi mà nhiều người bán và doanh nghiệp quan tâm. ")
        relationship_story.add_run("Dữ liệu có thể giúp trả lời câu hỏi này:")
        
        for run in relationship_story.runs:
            set_vietnamese_font(run)
            run.font.size = Pt(12)
        
        # Tìm khoảng giá có doanh số cao nhất
        max_sales_range = max(stats["price_sales_relationship"], key=lambda x: x.get("avg_sales", 0) or 0)
        if max_sales_range.get("avg_sales"):
            insight_story = doc.add_paragraph()
            insight_story.add_run("Khoảng giá ")
            insight_story.add_run(f'"{max_sales_range["price_range"]}"').bold = True
            insight_story.add_run(" có doanh số trung bình cao nhất. ")
            insight_story.add_run("Điều này cho thấy đây là 'vùng giá vàng' - mức giá mà người tiêu dùng cảm thấy hợp lý và sẵn sàng mua nhất. ")
            insight_story.add_run("Đây là insight quý giá cho các doanh nghiệp khi định giá sản phẩm.")
            
            for run in insight_story.runs:
                set_vietnamese_font(run)
                run.font.size = Pt(12)
        
        doc.add_paragraph()  # Spacing
    
    # 7. Câu chuyện về khuyến mãi
    if stats.get("discount_impact"):
        doc.add_heading("7. Câu Chuyện: Giảm Giá Có Thực Sự Giúp Bán Được Nhiều Hơn?", 1)
        discount_story = doc.add_paragraph()
        discount_story.add_run("Các chương trình giảm giá và khuyến mãi là công cụ marketing phổ biến trên Tiki. ")
        discount_story.add_run("Một câu hỏi nghiên cứu: Liệu giảm giá có thực sự ảnh hưởng đến doanh số không? ")
        discount_story.add_run("Dữ liệu có thể cung cấp câu trả lời:")
        
        for run in discount_story.runs:
            set_vietnamese_font(run)
            run.font.size = Pt(12)
        
        # So sánh sản phẩm có và không có discount
        no_discount = next((r for r in stats["discount_impact"] if "Không giảm giá" in r["discount_range"]), None)
        with_discount = next((r for r in stats["discount_impact"] if "Không giảm giá" not in r["discount_range"]), None)
        
        if no_discount and with_discount:
            comparison_story = doc.add_paragraph()
            no_discount_sales = no_discount.get("avg_sales") or 0
            with_discount_sales = with_discount.get("avg_sales") or 0
            
            if with_discount_sales > no_discount_sales:
                diff = ((with_discount_sales - no_discount_sales) / no_discount_sales) * 100 if no_discount_sales > 0 else 0
                comparison_story.add_run("Dữ liệu cho thấy sản phẩm có giảm giá thường bán được nhiều hơn. ")
                comparison_story.add_run("Điều này chứng minh rằng khuyến mãi là một công cụ marketing hiệu quả để thu hút người mua và tăng doanh số.")
            else:
                comparison_story.add_run("Thú vị là, một số sản phẩm không giảm giá vẫn bán rất chạy. ")
                comparison_story.add_run("Điều này có thể do chất lượng sản phẩm tốt, thương hiệu mạnh, hoặc đáp ứng đúng nhu cầu của người mua.")
            
            for run in comparison_story.runs:
                set_vietnamese_font(run)
                run.font.size = Pt(12)
        
        doc.add_paragraph()  # Spacing
    
    # 8. Kết luận
    doc.add_heading("8. Kết Luận: Những Câu Chuyện Đã Kể", 1)
    
    # Tổng hợp câu chuyện
    conclusion = doc.add_paragraph()
    conclusion.add_run("Qua quá trình phân tích dữ liệu, đã khám phá được nhiều insights thú vị về thị trường thương mại điện tử Việt Nam. ")
    conclusion.add_run(f"Với ")
    conclusion.add_run(f"{overview['total_products']:,}").bold = True
    conclusion.add_run(" sản phẩm từ ")
    conclusion.add_run(f"{overview['total_categories']:,}").bold = True
    conclusion.add_run(" danh mục, dataset này phản ánh một thị trường sôi động, đa dạng và đầy tiềm năng.")
    
    for run in conclusion.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # Những phát hiện chính
    doc.add_heading("Những Phát Hiện Chính", 2)
    learnings_intro = doc.add_paragraph()
    learnings_intro.add_run("Từ phân tích dữ liệu, có thể rút ra những phát hiện sau:")
    
    for run in learnings_intro.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(12)
    
    # Tính toán learnings từ dữ liệu
    learnings = []
    
    if prod_stats.get("distinct_brands"):
        learnings.append(f"Thị trường có {prod_stats['distinct_brands']:,} thương hiệu khác nhau, cho thấy sự đa dạng và cạnh tranh trên thị trường.")
    
    if prod_stats.get("discounted_products") and overview.get("total_products"):
        discount_ratio = (prod_stats["discounted_products"] / overview["total_products"]) * 100
        learnings.append(f"Khuyến mãi là một chiến lược phổ biến - có {discount_ratio:.1f}% sản phẩm đang giảm giá.")
    
    if prod_stats.get("avg_rating"):
        learnings.append(f"Người tiêu dùng khá hài lòng với sản phẩm trên thị trường, với điểm đánh giá trung bình {safe_format(prod_stats.get('avg_rating'), '.2f')}/5.0.")
    
    if stats.get("price_sales_relationship"):
        max_sales_range = max(stats["price_sales_relationship"], key=lambda x: x.get("avg_sales", 0) or 0)
        if max_sales_range.get("price_range"):
            learnings.append(f"Có một 'vùng giá vàng' - khoảng giá '{max_sales_range['price_range']}' nơi sản phẩm bán chạy nhất.")
    
    if cat_stats.get("distinct_levels"):
        learnings.append(f"Cấu trúc danh mục được tổ chức rất chặt chẽ với {cat_stats['distinct_levels']} cấp độ, giúp người mua dễ dàng tìm kiếm.")
    
    for learning in learnings:
        p = doc.add_paragraph(learning, style="List Bullet")
        for run in p.runs:
            set_vietnamese_font(run)
            run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # Ứng dụng
    doc.add_heading("Ứng Dụng Của Dataset", 2)
    application = doc.add_paragraph()
    application.add_run("Dataset này có thể được sử dụng cho nhiều mục đích nghiên cứu và phân tích:")
    
    for run in application.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(12)
    
    applications = [
        "Nghiên cứu về hành vi mua sắm của người Việt Nam",
        "Phân tích xu hướng thị trường và dự đoán tương lai",
        "So sánh các thương hiệu và sản phẩm",
        "Học tập về Data Engineering và Data Analytics với dữ liệu thực tế",
        "Phát triển các dự án phân tích dữ liệu khác"
    ]
    
    for app in applications:
        p = doc.add_paragraph(app, style="List Bullet")
        for run in p.runs:
            set_vietnamese_font(run)
            run.font.size = Pt(12)
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph("---")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = doc.add_paragraph("Tiki Data Pipeline - Data Story Document")
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_text.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Lưu file (xóa file cũ nếu tồn tại và đang bị lock)
    try:
        if output_path.exists():
            try:
                output_path.unlink()
            except PermissionError:
                print(f"⚠️  File đang được mở, vui lòng đóng file {output_path} và chạy lại script")
                raise
        doc.save(str(output_path))
        print(f"✅ Đã tạo file docx: {output_path}")
    except PermissionError as e:
        print(f"❌ Không thể ghi file: {e}")
        print(f"💡 Vui lòng đóng file {output_path} nếu đang mở và chạy lại script")
        raise


def main():
    """Hàm chính"""
    # Set UTF-8 encoding for Windows
    if sys.platform == "win32":
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8")
    
    print("=" * 70)
    print("📖 XÂY DỰNG CÂU CHUYỆN DỮ LIỆU - TIKI DATA PIPELINE")
    print("=" * 70)
    
    # Load config
    config = load_env_config()
    print(f"\n📋 Thông tin kết nối:")
    print(f"   Host: {config['host']}")
    print(f"   Port: {config['port']}")
    print(f"   Database: {config['database']}")
    print(f"   User: {config['user']}")
    
    # Kết nối database
    conn = connect_database(config)
    
    try:
        # Phân tích dữ liệu
        print("\n🔍 Bắt đầu phân tích dữ liệu...")
        stats = analyze_database(conn)
        
        # Tạo file docx (luôn ghi đè vào cùng một file)
        output_dir = PROJECT_ROOT / "docs"
        output_dir.mkdir(exist_ok=True)
        output_path = output_dir / "data_story.docx"
        
        print("\n📝 Đang tạo file docx...")
        create_document(stats, output_path)
        
        print("\n" + "=" * 70)
        print("✅ HOÀN TẤT!")
        print("=" * 70)
        print(f"\n📄 File đã được tạo tại: {output_path}")
        if output_path.exists():
            print(f"   Kích thước: {output_path.stat().st_size / 1024:.2f} KB")
        
    except Exception as e:
        print(f"\n❌ Lỗi: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
    finally:
        conn.close()
        print("\n🔌 Đã đóng kết nối database")


if __name__ == "__main__":
    main()

