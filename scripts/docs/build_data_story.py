"""
Script để phân tích database và xây dựng câu chuyện dữ liệu (Data Story)
Tạo file docx chứa câu chuyện dữ liệu của dự án Tiki Data Pipeline
"""

import os
import sys
import argparse
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

try:
    import psycopg2
    from psycopg2.extras import RealDictCursor
except ImportError:
    print("❌ Cần cài đặt psycopg2-binary: pip install psycopg2-binary")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("❌ Cần cài đặt python-docx: pip install python-docx")
    sys.exit(1)

# Google Drive API (optional)
try:
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import InstalledAppFlow
    from google.auth.transport.requests import Request
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload
    import pickle
    GOOGLE_DRIVE_AVAILABLE = True
except ImportError:
    GOOGLE_DRIVE_AVAILABLE = False

# Thêm src vào path để import modules
PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT / "src"))

# Đọc thông tin từ .env
def load_env_config():
    """Đọc cấu hình từ file .env"""
    env_file = PROJECT_ROOT / ".env"
    config = {
        "host": "localhost",
        "port": 5432,
        "database": "crawl_data",
        "user": "postgres",
        "password": "postgres",
    }
    
    if env_file.exists():
        with open(env_file, encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith("#") and "=" in line:
                    key, value = line.split("=", 1)
                    key = key.strip()
                    value = value.strip().strip('"').strip("'")
                    
                    if key == "POSTGRES_HOST":
                        config["host"] = value
                    elif key == "POSTGRES_PORT":
                        config["port"] = int(value)
                    elif key == "POSTGRES_USER":
                        config["user"] = value
                    elif key == "POSTGRES_PASSWORD":
                        config["password"] = value
                    elif key == "POSTGRES_DB" and value == "crawl_data":
                        config["database"] = value
    
    # Override với environment variables nếu có
    config["host"] = os.getenv("POSTGRES_HOST", config["host"])
    config["port"] = int(os.getenv("POSTGRES_PORT", config["port"]))
    config["user"] = os.getenv("POSTGRES_USER", config["user"])
    config["password"] = os.getenv("POSTGRES_PASSWORD", config["password"])
    config["database"] = os.getenv("POSTGRES_DB", config["database"])
    
    return config


def connect_database(config: dict[str, Any]):
    """Kết nối đến PostgreSQL database"""
    try:
        conn = psycopg2.connect(
            host=config["host"],
            port=config["port"],
            database=config["database"],
            user=config["user"],
            password=config["password"],
            connect_timeout=10,
        )
        print(f"✅ Đã kết nối đến database: {config['database']}")
        return conn
    except Exception as e:
        print(f"❌ Lỗi kết nối database: {e}")
        sys.exit(1)


def analyze_database(conn) -> dict[str, Any]:
    """Phân tích dữ liệu trong database"""
    stats = {}
    
    with conn.cursor(cursor_factory=RealDictCursor) as cur:
        # 1. Tổng quan
        print("📊 Đang phân tích tổng quan...")
        cur.execute("""
            SELECT 
                (SELECT COUNT(*) FROM categories) as total_categories,
                (SELECT COUNT(*) FROM products) as total_products,
                (SELECT COUNT(DISTINCT category_url) FROM products WHERE category_url IS NOT NULL) as categories_with_products,
                (SELECT COUNT(DISTINCT brand) FROM products WHERE brand IS NOT NULL) as total_brands,
                (SELECT COUNT(CASE WHEN sales_count > 0 THEN 1 END) FROM products WHERE sales_count IS NOT NULL) as products_sold,
                (SELECT COUNT(CASE WHEN sales_count > 1000 THEN 1 END) FROM products WHERE sales_count IS NOT NULL) as bestsellers
        """)
        stats["overview"] = dict(cur.fetchone())
        
        # 2. Thống kê categories
        print("📁 Đang phân tích categories...")
        cur.execute("""
            SELECT 
                COUNT(*) as total,
                COUNT(DISTINCT level) as distinct_levels,
                MIN(level) as min_level,
                MAX(level) as max_level,
                AVG(product_count) as avg_products_per_category,
                SUM(product_count) as total_product_count_in_categories
            FROM categories
        """)
        stats["categories"] = dict(cur.fetchone())
        
        # 3. Thống kê products
        print("🛍️ Đang phân tích products...")
        cur.execute("""
            SELECT 
                COUNT(*) as total,
                COUNT(DISTINCT category_url) as categories_covered,
                COUNT(DISTINCT brand) as distinct_brands,
                COUNT(DISTINCT seller_id) as distinct_sellers,
                COUNT(CASE WHEN seller_is_official = TRUE THEN 1 END) as official_seller_products,
                AVG(sales_count) as avg_sales_count,
                MAX(sales_count) as max_sales_count,
                MIN(sales_count) as min_sales_count,
                AVG(price) as avg_price,
                MAX(price) as max_price,
                MIN(price) as min_price,
                AVG(rating_average) as avg_rating,
                AVG(review_count) as avg_reviews,
                COUNT(CASE WHEN stock_available = TRUE THEN 1 END) as in_stock_count,
                COUNT(CASE WHEN discount_percent > 0 THEN 1 END) as discounted_products,
                AVG(discount_percent) as avg_discount_percent
            FROM products
        """)
        stats["products"] = dict(cur.fetchone())
        
        # 4. Top categories theo số lượng sản phẩm
        print("🏆 Đang lấy top categories...")
        cur.execute("""
            SELECT 
                c.name,
                c.level,
                COUNT(p.id) as product_count,
                AVG(p.price) as avg_price,
                AVG(p.sales_count) as avg_sales
            FROM categories c
            LEFT JOIN products p ON c.url = p.category_url
            GROUP BY c.id, c.name, c.level
            HAVING COUNT(p.id) > 0
            ORDER BY product_count DESC
            LIMIT 10
        """)
        stats["top_categories"] = [dict(row) for row in cur.fetchall()]
        
        # 5. Top products theo sales_count
        print("⭐ Đang lấy top products...")
        cur.execute("""
            SELECT 
                product_id,
                name,
                price,
                sales_count,
                rating_average,
                review_count,
                brand,
                seller_name
            FROM products
            WHERE sales_count IS NOT NULL
            ORDER BY sales_count DESC
            LIMIT 10
        """)
        stats["top_products"] = [dict(row) for row in cur.fetchall()]
        
        # 6. Phân bố giá
        print("💰 Đang phân tích phân bố giá...")
        cur.execute("""
            SELECT 
                CASE 
                    WHEN price < 100000 THEN 'Dưới 100k'
                    WHEN price < 500000 THEN '100k - 500k'
                    WHEN price < 1000000 THEN '500k - 1M'
                    WHEN price < 5000000 THEN '1M - 5M'
                    ELSE 'Trên 5M'
                END as price_range,
                COUNT(*) as count,
                AVG(price) as avg_price
            FROM products
            WHERE price IS NOT NULL
            GROUP BY price_range
            ORDER BY MIN(price)
        """)
        stats["price_distribution"] = [dict(row) for row in cur.fetchall()]
        
        # 7. Phân bố rating
        print("⭐ Đang phân tích rating...")
        cur.execute("""
            SELECT 
                CASE 
                    WHEN rating_average IS NULL THEN 'Chưa có rating'
                    WHEN rating_average < 3.0 THEN 'Dưới 3.0'
                    WHEN rating_average < 4.0 THEN '3.0 - 4.0'
                    WHEN rating_average < 4.5 THEN '4.0 - 4.5'
                    ELSE 'Trên 4.5'
                END as rating_range,
                COUNT(*) as count
            FROM products
            GROUP BY rating_range
            ORDER BY MIN(COALESCE(rating_average, 0))
        """)
        stats["rating_distribution"] = [dict(row) for row in cur.fetchall()]
        
        # 8. Thống kê theo brand
        print("🏷️ Đang phân tích brands...")
        cur.execute("""
            SELECT 
                brand,
                COUNT(*) as product_count,
                AVG(price) as avg_price,
                AVG(sales_count) as avg_sales,
                AVG(rating_average) as avg_rating
            FROM products
            WHERE brand IS NOT NULL
            GROUP BY brand
            ORDER BY product_count DESC
            LIMIT 10
        """)
        stats["top_brands"] = [dict(row) for row in cur.fetchall()]
        
        # 9. Phân tích computed fields
        print("📈 Đang phân tích computed fields...")
        cur.execute("""
            SELECT 
                AVG(estimated_revenue) as avg_revenue,
                SUM(estimated_revenue) as total_revenue,
                AVG(popularity_score) as avg_popularity,
                MAX(popularity_score) as max_popularity,
                AVG(value_score) as avg_value_score,
                MAX(value_score) as max_value_score,
                COUNT(CASE WHEN price_category = 'budget' THEN 1 END) as budget_count,
                COUNT(CASE WHEN price_category = 'mid-range' THEN 1 END) as midrange_count,
                COUNT(CASE WHEN price_category = 'premium' THEN 1 END) as premium_count,
                COUNT(CASE WHEN price_category = 'luxury' THEN 1 END) as luxury_count
            FROM products
            WHERE estimated_revenue IS NOT NULL
        """)
        stats["computed_fields"] = dict(cur.fetchone())
        
        # 10. Mối quan hệ giữa giá và doanh số
        print("🔗 Đang phân tích mối quan hệ giá-doanh số...")
        cur.execute("""
            SELECT 
                CASE 
                    WHEN price < 100000 THEN 'Dưới 100k'
                    WHEN price < 500000 THEN '100k - 500k'
                    WHEN price < 1000000 THEN '500k - 1M'
                    WHEN price < 5000000 THEN '1M - 5M'
                    ELSE 'Trên 5M'
                END as price_range,
                AVG(sales_count) as avg_sales,
                AVG(rating_average) as avg_rating,
                COUNT(*) as product_count
            FROM products
            WHERE price IS NOT NULL AND sales_count IS NOT NULL
            GROUP BY price_range
            ORDER BY MIN(price)
        """)
        stats["price_sales_relationship"] = [dict(row) for row in cur.fetchall()]
        
        # 11. Mối quan hệ giữa discount và sales (loại bỏ truy vấn lặp)
        print("💰 Đang phân tích tác động của discount...")
        cur.execute("""
            SELECT 
                CASE 
                    WHEN discount_percent = 0 OR discount_percent IS NULL THEN 'Không giảm giá'
                    WHEN discount_percent < 10 THEN 'Giảm dưới 10%'
                    WHEN discount_percent < 20 THEN 'Giảm 10-20%'
                    WHEN discount_percent < 30 THEN 'Giảm 20-30%'
                    ELSE 'Giảm trên 30%'
                END as discount_range,
                AVG(sales_count) as avg_sales,
                COUNT(*) as product_count,
                AVG(rating_average) as avg_rating
            FROM products
            WHERE sales_count IS NOT NULL
            GROUP BY discount_range
            ORDER BY MIN(COALESCE(discount_percent, 0))
        """)
        stats["discount_impact"] = [dict(row) for row in cur.fetchall()]
        
        # 12. Top products theo computed fields
        print("⭐ Đang lấy top products theo computed fields...")
        cur.execute("""
            SELECT 
                product_id,
                name,
                price,
                sales_count,
                rating_average,
                popularity_score,
                value_score,
                estimated_revenue,
                discount_percent
            FROM products
            WHERE popularity_score IS NOT NULL
            ORDER BY popularity_score DESC
            LIMIT 5
        """)
        stats["top_by_popularity"] = [dict(row) for row in cur.fetchall()]
        
        cur.execute("""
            SELECT 
                product_id,
                name,
                price,
                sales_count,
                rating_average,
                popularity_score,
                value_score,
                estimated_revenue
            FROM products
            WHERE value_score IS NOT NULL
            ORDER BY value_score DESC
            LIMIT 5
        """)
        stats["top_by_value"] = [dict(row) for row in cur.fetchall()]
        
        # 13. Phân tích sự phân bố sales
        print("📈 Đang phân tích phân bố doanh số...")
        cur.execute("""
            SELECT 
                COUNT(CASE WHEN sales_count = 0 THEN 1 END) as no_sales,
                COUNT(CASE WHEN sales_count > 0 AND sales_count <= 100 THEN 1 END) as low_sales,
                COUNT(CASE WHEN sales_count > 100 AND sales_count <= 500 THEN 1 END) as medium_sales,
                COUNT(CASE WHEN sales_count > 500 AND sales_count <= 1000 THEN 1 END) as high_sales,
                COUNT(CASE WHEN sales_count > 1000 THEN 1 END) as bestsellers,
                MAX(sales_count) as max_sales,
                PERCENTILE_CONT(0.5) WITHIN GROUP(ORDER BY sales_count) as median_sales,
                AVG(sales_count) as avg_sales
            FROM products
            WHERE sales_count IS NOT NULL
        """)
        stats["sales_distribution"] = dict(cur.fetchone())
        
        # 14. Top 5 categories từ products
        print("🏆 Đang lấy top categories theo doanh số...")
        cur.execute("""
            SELECT 
                category_url,
                COUNT(*) as product_count,
                SUM(sales_count) as total_sales,
                AVG(price) as avg_price,
                AVG(rating_average) as avg_rating,
                AVG(discount_percent) as avg_discount
            FROM products
            WHERE category_url IS NOT NULL
            GROUP BY category_url
            ORDER BY total_sales DESC
            LIMIT 10
        """)
        stats["top_categories_by_sales"] = [dict(row) for row in cur.fetchall()]
        
        # 15. Phân tích official sellers
        print("👑 Đang phân tích official sellers...")
        cur.execute("""
            SELECT 
                COUNT(CASE WHEN seller_is_official = TRUE THEN 1 END) as official_products,
                COUNT(CASE WHEN seller_is_official = FALSE THEN 1 END) as third_party_products,
                AVG(CASE WHEN seller_is_official = TRUE THEN sales_count END) as avg_sales_official,
                AVG(CASE WHEN seller_is_official = FALSE THEN sales_count END) as avg_sales_third_party,
                AVG(CASE WHEN seller_is_official = TRUE THEN rating_average END) as avg_rating_official,
                AVG(CASE WHEN seller_is_official = FALSE THEN rating_average END) as avg_rating_third_party
            FROM products
        """)
        stats["official_analysis"] = dict(cur.fetchone())
        
        # 16. Phân tích stock
        print("📦 Đang phân tích stock...")
        cur.execute("""
            SELECT 
                COUNT(CASE WHEN stock_available = TRUE THEN 1 END) as in_stock,
                COUNT(CASE WHEN stock_available = FALSE THEN 1 END) as out_of_stock,
                AVG(CASE WHEN stock_available = TRUE THEN sales_count END) as avg_sales_in_stock,
                AVG(CASE WHEN stock_available = FALSE THEN sales_count END) as avg_sales_out_of_stock
            FROM products
        """)
        stats["stock_analysis"] = dict(cur.fetchone())
    
    return stats


def safe_format(value: Any, format_str: str = ",.0f") -> str:
    """Format số an toàn, xử lý None"""
    if value is None:
        return "N/A"
    try:
        return f"{value:{format_str}}"
    except (ValueError, TypeError):
        return str(value) if value else "N/A"


def upload_to_google_drive(file_path: Path, folder_id: Optional[str] = None) -> Optional[str]:
    """Upload file lên Google Drive
    
    Args:
        file_path: Đường dẫn file cần upload
        folder_id: ID của folder trên Google Drive (optional)
    
    Returns:
        File ID nếu upload thành công, None nếu thất bại
    """
    if not GOOGLE_DRIVE_AVAILABLE:
        print("\n❌ Google Drive API không khả dụng")
        print("💡 Cài đặt: pip install google-auth google-auth-oauthlib google-auth-httplib2 google-api-python-client")
        return None
    
    # Đường dẫn file credentials
    credentials_file = PROJECT_ROOT /"docs"/ "credentials" / "google_drive_credentials.json"
    token_file = PROJECT_ROOT / "docs" / "credentials" / "token.pickle"
    
    if not credentials_file.exists():
        print(f"\n❌ Không tìm thấy credentials file: {credentials_file}")
        print("💡 Tạo credentials tại: https://console.cloud.google.com/apis/credentials")
        print("   - Tạo OAuth 2.0 Client ID (Desktop app)")
        print("   - Download JSON và lưu vào credentials/google_drive_credentials.json")
        return None
    
    creds = None
    
    # Load token nếu đã có
    if token_file.exists():
        with open(token_file, 'rb') as token:
            creds = pickle.load(token)
    
    # Nếu không có credentials hợp lệ, yêu cầu login
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            print("🔄 Đang refresh token...")
            creds.refresh(Request())
        else:
            print("🔐 Đang mở browser để xác thực Google Drive...")
            flow = InstalledAppFlow.from_client_secrets_file(
                str(credentials_file),
                scopes=['https://www.googleapis.com/auth/drive.file']
            )
            creds = flow.run_local_server(port=0)
        
        # Lưu token
        token_file.parent.mkdir(exist_ok=True)
        with open(token_file, 'wb') as token:
            pickle.dump(creds, token)
        print("✅ Đã lưu token")
    
    try:
        # Tạo service
        service = build('drive', 'v3', credentials=creds)
        
        # Metadata file
        file_metadata = {
            'name': file_path.name,
        }
        
        if folder_id:
            file_metadata['parents'] = [folder_id]
        
        # Kiểm tra xem file đã tồn tại chưa
        query = f"name='{file_path.name}'"
        if folder_id:
            query += f" and '{folder_id}' in parents"
        query += " and trashed=false"
        
        results = service.files().list(
            q=query,
            spaces='drive',
            fields='files(id, name)'
        ).execute()
        
        existing_files = results.get('files', [])
        
        media = MediaFileUpload(str(file_path), mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        
        if existing_files:
            # Update file hiện tại
            file_id = existing_files[0]['id']
            print(f"\n🔄 Đang cập nhật file trên Google Drive...")
            file = service.files().update(
                fileId=file_id,
                media_body=media
            ).execute()
            print(f"✅ Đã cập nhật file: {file.get('name')}")
        else:
            # Tạo file mới
            print(f"\n📤 Đang upload file lên Google Drive...")
            file = service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id, name, webViewLink'
            ).execute()
            print(f"✅ Đã upload file: {file.get('name')}")
        
        # Lấy link
        file_id = file.get('id')
        file_link = service.files().get(
            fileId=file_id,
            fields='webViewLink'
        ).execute()
        
        print(f"🔗 Link: {file_link.get('webViewLink')}")
        return file_id
        
    except Exception as e:
        print(f"\n❌ Lỗi khi upload lên Google Drive: {e}")
        import traceback
        traceback.print_exc()
        return None


def create_document(stats: dict[str, Any], output_path: Path):
    """Tạo file docx với câu chuyện dữ liệu"""
    doc = Document()
    
    # Cấu hình font cho tiếng Việt
    def set_vietnamese_font(run):
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    
    # Helper: tạo tiêu đề mục dạng bullet (indent 0.25")
    def add_section_title(text: str):
        p = doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            set_vietnamese_font(run)
            run.bold = True
            run.font.size = Pt(12)
        return p
    
    # Helper: định dạng đoạn nội dung thành sub-bullet (indent 0.5")
    def format_as_subbullet(paragraph):
        paragraph.style = doc.styles["List Bullet 2"]
        paragraph.paragraph_format.left_indent = Inches(0.5)
        for run in paragraph.runs:
            set_vietnamese_font(run)
            if run.font.size is None:
                run.font.size = Pt(12)
    
    # Helper: tạo sub-bullet nhanh từ text
    def add_subbullet_text(text: str, bold: bool = False):
        p = doc.add_paragraph(style="List Bullet 2")
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(text)
        set_vietnamese_font(run)
        run.font.size = Pt(12)
        if bold:
            run.bold = True
        return p
    
    # Helper: spacer dòng trống có kiểm soát
    def add_spacer(lines: int = 1):
        for _ in range(lines):
            doc.add_paragraph()
    
    # Title
    title = doc.add_heading("Câu Chuyện Dữ Liệu - Tiki Data Pipeline", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Subtitle
    subtitle = doc.add_paragraph(f"Ngày cập nhật: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()  # Spacing
    
    # ============================================
    # PHẦN BỐI CẢNH VÀ GIỚI THIỆU
    # ============================================
    
    doc.add_heading("Bối Cảnh: Thị Trường Thương Mại Điện Tử Việt Nam", 1)
    
    # Bối cảnh thị trường
    add_subbullet_text(
        "Thương mại điện tử Việt Nam đang phát triển mạnh mẽ, dự kiến đạt 49 tỷ USD vào năm 2025. Tiki.vn là một trong những nền tảng thương mại điện tử hàng đầu, được thành lập từ năm 2010 với hàng triệu sản phẩm đa dạng. Dữ liệu từ Tiki.vn phản ánh xu hướng mua sắm, hành vi tiêu dùng và cấu trúc thị trường, có giá trị nghiên cứu cao."
    )
    
    doc.add_paragraph()  # Spacing
    
    # Ý nghĩa của dataset
    add_subbullet_text(
        "Dataset này không chỉ là danh sách sản phẩm, mà là cửa sổ để hiểu về thị trường thương mại điện tử Việt Nam. Từ dataset có thể khám phá: xu hướng tiêu dùng, cấu trúc thị trường, hành vi mua sắm, sự cạnh tranh giữa các thương hiệu, và giá trị thị trường."
    )
    
    doc.add_paragraph()  # Spacing
    
    # Lý do chọn đề tài
    doc.add_heading("Lý Do Chọn Đề Tài", 1)
    
    add_subbullet_text("Việc xây dựng dataset và phân tích dữ liệu từ Tiki.vn được lựa chọn dựa trên những lý do sau:")
    
    reasons = [
        ("Tầm quan trọng của thương mại điện tử", 
         "Thương mại điện tử đang là xu hướng tất yếu của thời đại số. Việc hiểu rõ thị trường này không chỉ quan trọng đối với các doanh nghiệp, mà còn có giá trị nghiên cứu và học thuật cao. Dữ liệu từ Tiki.vn cung cấp cái nhìn toàn diện về hành vi tiêu dùng, xu hướng thị trường và cấu trúc kinh doanh."),
        
        ("Giá trị thực tiễn của dữ liệu", 
         "Khác với các dataset mẫu trong sách giáo khoa, dữ liệu từ Tiki.vn là dữ liệu thực tế từ thị trường. Dataset này có thể được sử dụng để nghiên cứu thị trường, phân tích cạnh tranh, dự đoán xu hướng. Đây là cơ hội để làm việc với dữ liệu thực tế và áp dụng kiến thức đã học."),
        
        ("Thách thức kỹ thuật và cơ hội học hỏi", 
         "Xây dựng hệ thống thu thập dữ liệu từ website động như Tiki.vn là một thách thức kỹ thuật. Dự án này đòi hỏi kiến thức về web scraping, xử lý dữ liệu, thiết kế database, và phân tích dữ liệu. Đây là cơ hội để áp dụng và nâng cao các kỹ năng trong lĩnh vực Data Engineering và Data Analytics."),
        
        ("Tính mới và đóng góp", 
         "Mặc dù có nhiều nghiên cứu về thương mại điện tử, nhưng việc xây dựng một dataset toàn diện và có hệ thống từ Tiki.vn vẫn còn hạn chế. Dataset này có thể đóng góp cho cộng đồng nghiên cứu, các nhà phân tích dữ liệu, và những người quan tâm đến thị trường thương mại điện tử Việt Nam."),
        
        ("Ứng dụng trong giáo dục và nghiên cứu", 
         "Dataset này có thể được sử dụng như một case study trong giáo dục về Data Engineering, Data Analytics, và Business Intelligence. Nó cung cấp một ví dụ thực tế về cách thu thập, xử lý, và phân tích dữ liệu từ nguồn thực tế, giúp sinh viên và nhà nghiên cứu hiểu rõ hơn về quy trình làm việc với dữ liệu thực tế."),
        
        ("Tiềm năng mở rộng", 
         "Dự án này có tiềm năng mở rộng lớn. Có thể phát triển thành một hệ thống monitoring thị trường, một công cụ phân tích cạnh tranh, hoặc một nền tảng cung cấp dữ liệu cho các ứng dụng khác. Dataset có thể được cập nhật định kỳ để theo dõi sự thay đổi của thị trường.")
    ]
    
    for idx, (title, content) in enumerate(reasons, 1):
        reason_heading = doc.add_paragraph(f"{idx}. {title}", style="List Bullet")
        reason_heading.paragraph_format.left_indent = Inches(0.25)
        for run in reason_heading.runs:
            set_vietnamese_font(run)
            run.bold = True
            run.font.size = Pt(12)
        
        reason_para = doc.add_paragraph(content, style="List Bullet 2")
        reason_para.paragraph_format.left_indent = Inches(0.5)
        for run in reason_para.runs:
            set_vietnamese_font(run)
            run.font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # Lời mở đầu
    doc.add_heading("Lời Mở Đầu: Câu Chuyện Từ Dữ Liệu", 1)
    add_subbullet_text(
        "Đằng sau mỗi con số là một câu chuyện. Đằng sau mỗi sản phẩm là một lựa chọn của người tiêu dùng. Đằng sau mỗi danh mục là một xu hướng thị trường. Tài liệu này trình bày những câu chuyện được khám phá từ dữ liệu thu thập từ Tiki.vn, một trong những nền tảng thương mại điện tử hàng đầu Việt Nam."
    )
    
    doc.add_paragraph()  # Spacing
    
    # Giới thiệu về dataset
    doc.add_heading("Về Dataset", 2)
    add_subbullet_text(
        "Dataset này chứa thông tin về hàng nghìn sản phẩm từ Tiki.vn, được thu thập và xử lý một cách có hệ thống. Mỗi sản phẩm trong dataset bao gồm thông tin chi tiết về: tên sản phẩm, giá cả, mô tả, đánh giá của người dùng, thông tin người bán, thương hiệu, số lượng đã bán, và nhiều chỉ số phân tích khác. Mỗi dòng dữ liệu không chỉ phản ánh thông tin về sản phẩm, mà còn cho thấy về thị trường, về hành vi mua sắm của người tiêu dùng, và về những xu hướng đang diễn ra."
    )
    
    doc.add_paragraph()  # Spacing
    
    # Câu chuyện từ dữ liệu
    doc.add_heading("Những Câu Hỏi Nghiên Cứu", 2)
    add_subbullet_text(
        "Khi bắt đầu với dataset này, có nhiều câu hỏi nghiên cứu được đặt ra. Dữ liệu sẽ giúp trả lời những câu hỏi đó. Dưới đây là những vấn đề có thể khám phá từ dataset:"
    )
    
    story_points = [
        "Thị trường Tiki có quy mô như thế nào? Có bao nhiêu sản phẩm và danh mục?",
        "Người tiêu dùng đang mua gì? Sản phẩm nào được mua nhiều nhất?",
        "Giá cả trên thị trường phân bố như thế nào? Người tiêu dùng thường mua ở mức giá nào?",
        "Thương hiệu nào đang dẫn đầu? Ai là người bán tốt nhất?",
        "Người tiêu dùng đánh giá sản phẩm như thế nào? Điểm số và review phản ánh điều gì?",
        "Giảm giá có thực sự ảnh hưởng đến doanh số không?"
    ]
    
    for point in story_points:
        p = doc.add_paragraph(point, style="List Bullet 2")
        for run in p.runs:
            set_vietnamese_font(run)
            run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # 1. Tổng quan (bullet format)
    add_section_title("1. Tổng Quan Dữ Liệu: Mẫu Nghiên Cứu")
    overview = stats["overview"]
    overview_text = doc.add_paragraph()
    overview_text.add_run("Sau quá trình thu thập và xử lý, dataset đã được xây dựng với quy mô đáng kể. ")
    overview_text.add_run(f"Dataset hiện tại bao gồm ")
    overview_text.add_run(f"{overview['total_products']:,}").bold = True
    overview_text.add_run(" sản phẩm từ ")
    overview_text.add_run(f"{overview['categories_with_products']:,}").bold = True
    overview_text.add_run(" danh mục khác nhau, được cung cấp bởi ")
    overview_text.add_run(f"{overview['total_brands']:,}").bold = True
    overview_text.add_run(" thương hiệu. ")
    
    # Tính toán tỷ lệ bestsellers
    if overview['total_products'] > 0:
        bestseller_ratio = (overview['bestsellers'] / overview['total_products']) * 100
        overview_text.add_run(f"Thú vị là, có ")
        overview_text.add_run(f"{bestseller_ratio:.1f}%").bold = True
        overview_text.add_run(f" sản phẩm được xem là 'bestseller' (bán > 1000 cái).")
    
    format_as_subbullet(overview_text)
    add_spacer()
    
    doc.add_paragraph()  # Spacing
    
    # 2. Categories
    add_section_title("2. Câu Chuyện Về Danh Mục: Cấu Trúc Thị Trường")
    cat_text = doc.add_paragraph()
    cat_text.add_run("Danh mục sản phẩm phản ánh cách tổ chức và phân loại thị trường. ")
    
    if overview.get("categories_with_products"):
        cat_text.add_run(f"Dữ liệu cho thấy sản phẩm được phân bố vào ")
        cat_text.add_run(f"{overview['categories_with_products']:,}").bold = True
        cat_text.add_run(" danh mục khác nhau. ")
    
    cat_text.add_run("Mỗi danh mục đại diện cho một phân khúc thị trường với những đặc điểm riêng về giá cả, thương hiệu, và hành vi mua sắm của người tiêu dùng.")
    
    format_as_subbullet(cat_text)
    
    # Insights về danh mục từ top_categories_by_sales
    if stats.get("top_categories_by_sales"):
        top_cat_insight = doc.add_paragraph()
        top_cat = stats["top_categories_by_sales"][0] if stats["top_categories_by_sales"] else None
        if top_cat:
            top_cat_insight.add_run("Danh mục có doanh số cao nhất là: ")
            top_cat_insight.add_run(f"{top_cat['product_count']:,} sản phẩm").bold = True
            if top_cat.get("avg_price"):
                top_cat_insight.add_run(f" với giá trung bình ")
                top_cat_insight.add_run(f"{safe_format(top_cat.get('avg_price'), ',.0f')} VND").bold = True
            
            if top_cat.get("total_sales") and top_cat.get("product_count"):
                avg_sales_per_product = top_cat["total_sales"] / top_cat["product_count"]
                top_cat_insight.add_run(f". Trung bình mỗi sản phẩm bán được ")
                top_cat_insight.add_run(f"{avg_sales_per_product:.0f} cái").bold = True
            top_cat_insight.add_run(".")
            format_as_subbullet(top_cat_insight)
        add_spacer()
    
    doc.add_paragraph()  # Spacing
    
    # 3. Products
    add_section_title("3. Câu Chuyện Về Sản Phẩm: Thị Trường Trong Lòng Bàn Tay")
    prod_stats = stats["products"]
    prod_text = doc.add_paragraph()
    prod_text.add_run("Mỗi sản phẩm trong dataset phản ánh một lựa chọn của người tiêu dùng dựa trên nhu cầu, giá cả, và đánh giá. ")
    prod_text.add_run(f"Dataset bao gồm sản phẩm từ ")
    prod_text.add_run(f"{prod_stats['distinct_brands']:,}").bold = True
    prod_text.add_run(" thương hiệu khác nhau và ")
    prod_text.add_run(f"{prod_stats['distinct_sellers']:,}").bold = True
    prod_text.add_run(" người bán. ")
    prod_text.add_run("Điều này cho thấy thị trường rất đa dạng và cạnh tranh, tạo ra nhiều lựa chọn cho người tiêu dùng.")
    
    format_as_subbullet(prod_text)
    
    # Câu chuyện về giá cả
    if prod_stats.get("avg_price") is not None:
        price_story = doc.add_paragraph()
        price_story.add_run("Một câu hỏi nghiên cứu quan trọng: Người tiêu dùng thường mua ở mức giá nào? ")
        if stats["price_distribution"]:
            max_count_range = max(stats["price_distribution"], key=lambda x: x["count"])
            total_products = sum(r["count"] for r in stats["price_distribution"])
            max_percentage = (max_count_range["count"] / total_products) * 100 if total_products > 0 else 0
            price_story.add_run(f"Dữ liệu cho thấy ")
            price_story.add_run(f"{max_percentage:.1f}%").bold = True
            price_story.add_run(f" sản phẩm nằm trong khoảng ")
            price_story.add_run(f'"{max_count_range["price_range"]}"').bold = True
            price_story.add_run(". ")
            price_story.add_run("Điều này phản ánh phân khúc giá mà người tiêu dùng Việt Nam thường lựa chọn khi mua sắm online.")
        else:
            price_story.add_run(f"Giá trung bình là ")
            price_story.add_run(f"{safe_format(prod_stats.get('avg_price'), ',.0f')} VND").bold = True
            price_story.add_run(", cho thấy mức giá phổ biến trên thị trường.")
        
        format_as_subbullet(price_story)
    
    # Câu chuyện về sản phẩm bán chạy
    if stats["top_products"]:
        top_prod_story = doc.add_paragraph()
        top_prod = stats["top_products"][0] if stats["top_products"] else None
        if top_prod:
            top_prod_story.add_run("Sản phẩm nào được mua nhiều nhất? ")
            top_prod_story.add_run(f'"{top_prod["name"][:60]}"').bold = True
            if top_prod.get("sales_count"):
                top_prod_story.add_run(f" đã bán được ")
                top_prod_story.add_run(f"{safe_format(top_prod.get('sales_count'), ',')}").bold = True
                top_prod_story.add_run(" sản phẩm. ")
            top_prod_story.add_run("Điều này cho thấy sản phẩm này đáp ứng được nhu cầu và sở thích của đông đảo người tiêu dùng.")
            format_as_subbullet(top_prod_story)
        add_spacer()
    
    doc.add_paragraph()  # Spacing
    
    # 4. Brands
    if stats["top_brands"]:
        add_section_title("4. Câu Chuyện Về Thương Hiệu: Ai Đang Dẫn Đầu?")
        brand_text = doc.add_paragraph()
        brand_text.add_run("Thương hiệu đóng vai trò quan trọng trong quyết định mua sắm của người tiêu dùng. ")
        brand_text.add_run("Thương hiệu không chỉ là tên gọi, mà còn là lời hứa về chất lượng và giá trị. ")
        brand_text.add_run("Trên Tiki, có rất nhiều thương hiệu cạnh tranh với nhau để thu hút người mua, tạo nên một thị trường đa dạng và sôi động.")
        format_as_subbullet(brand_text)
        
        # Câu chuyện về thương hiệu hàng đầu
        top_brand = stats["top_brands"][0] if stats["top_brands"] else None
        if top_brand:
            brand_story = doc.add_paragraph()
            brand_story.add_run("Thương hiệu có nhiều sản phẩm nhất là ")
            brand_story.add_run(f'"{top_brand["brand"]}"').bold = True
            brand_story.add_run(f" với ")
            brand_story.add_run(f"{top_brand['product_count']:,}").bold = True
            brand_story.add_run(" sản phẩm. ")
            brand_story.add_run("Điều này cho thấy thương hiệu này đã xây dựng được một danh mục sản phẩm đa dạng và có vị thế mạnh trên thị trường.")
            format_as_subbullet(brand_story)
        add_spacer()
    
    doc.add_paragraph()  # Spacing
    
    # 5. Market value
    if stats.get("computed_fields"):
        add_section_title("5. Câu Chuyện Về Giá Trị Thị Trường")
        computed = stats["computed_fields"]
        
        # Estimated Revenue - câu chuyện về quy mô
        if computed.get("total_revenue"):
            revenue_story = doc.add_paragraph()
            revenue_story.add_run("Một câu hỏi nghiên cứu quan trọng: Thị trường này có giá trị bao nhiêu? ")
            revenue_story.add_run("Từ dataset, tổng doanh thu ước tính là ")
            revenue_story.add_run(f"{safe_format(computed.get('total_revenue') / 1000000000, '.2f')} tỷ VND").bold = True
            revenue_story.add_run(". ")
            revenue_story.add_run("Con số này phản ánh quy mô và tiềm năng của thị trường thương mại điện tử Việt Nam.")
            
            format_as_subbullet(revenue_story)
        
        doc.add_paragraph()  # Spacing
    
    # 6. Price-sales relationship
    if stats.get("price_sales_relationship"):
        add_section_title("6. Câu Chuyện: Giá Nào Bán Chạy Nhất?")
        relationship_story = doc.add_paragraph()
        relationship_story.add_run("Một câu hỏi nghiên cứu quan trọng: Ở mức giá nào thì sản phẩm bán chạy nhất? ")
        relationship_story.add_run("Đây là câu hỏi mà nhiều người bán và doanh nghiệp quan tâm. ")
        relationship_story.add_run("Dữ liệu có thể giúp trả lời câu hỏi này:")
        format_as_subbullet(relationship_story)
        
        # Tìm khoảng giá có doanh số cao nhất
        max_sales_range = max(stats["price_sales_relationship"], key=lambda x: x.get("avg_sales", 0) or 0)
        if max_sales_range.get("avg_sales"):
            insight_story = doc.add_paragraph()
            insight_story.add_run("Khoảng giá ")
            insight_story.add_run(f'"{max_sales_range["price_range"]}"').bold = True
            insight_story.add_run(" có doanh số trung bình cao nhất. ")
            insight_story.add_run("Điều này cho thấy đây là 'vùng giá vàng' - mức giá mà người tiêu dùng cảm thấy hợp lý và sẵn sàng mua nhất. ")
            insight_story.add_run("Đây là insight quý giá cho các doanh nghiệp khi định giá sản phẩm.")
            
            format_as_subbullet(insight_story)
        
        doc.add_paragraph()  # Spacing
    
    # 7. Sales distribution
    if stats.get("sales_distribution"):
        add_section_title("7. Câu Chuyện: Ai Là Những Sản Phẩm Bán Chạy?")
        sales_dist = stats["sales_distribution"]
        
        sales_story = doc.add_paragraph()
        sales_story.add_run("Không phải tất cả sản phẩm đều bán được như nhau. ")
        sales_story.add_run("Trên thị trường, chỉ có một tỉ lệ nhỏ sản phẩm được người tiêu dùng ưa chuộng. ")
        
        # Tính toán phân bố
        total_products = (sales_dist.get("no_sales") or 0) + (sales_dist.get("low_sales") or 0) + \
                        (sales_dist.get("medium_sales") or 0) + (sales_dist.get("high_sales") or 0) + \
                        (sales_dist.get("bestsellers") or 0)
        
        if total_products > 0:
            no_sales_pct = ((sales_dist.get("no_sales") or 0) / total_products) * 100
            bestseller_pct = ((sales_dist.get("bestsellers") or 0) / total_products) * 100
            
            sales_story.add_run(f"Dữ liệu cho thấy: ")
            sales_story.add_run(f"{no_sales_pct:.1f}%").bold = True
            sales_story.add_run(f" sản phẩm chưa bán được (có thể là sản phẩm mới hoặc chất lượng chưa tốt), ")
            sales_story.add_run(f"trong khi ")
            sales_story.add_run(f"{bestseller_pct:.1f}%").bold = True
            sales_story.add_run(f" là 'bestsellers' bán hơn 1000 cái. ")
            sales_story.add_run("Sự chênh lệch này cho thấy hành vi mua sắm rất tập trung vào một số sản phẩm 'sao' nhất định.")
        
        format_as_subbullet(sales_story)
        
        # Median vs Average insight
        if sales_dist.get("avg_sales") and sales_dist.get("median_sales"):
            median_insight = doc.add_paragraph()
            median_insight.add_run("Một phát hiện thú vị khác: ")
            median_insight.add_run(f"doanh số trung vị (median) là {safe_format(sales_dist.get('median_sales'), '.0f')} cái, ").bold = True
            median_insight.add_run(f"nhưng doanh số trung bình (average) là {safe_format(sales_dist.get('avg_sales'), '.0f')} cái. ")
            median_insight.add_run("Điều này cho thấy có một số sản phẩm bestseller 'kéo' doanh số trung bình lên rất cao. ")
            median_insight.add_run("Nói cách khác, thị trường có phân hóa lớn - có những sản phẩm bán cực chạy, nhưng nhiều sản phẩm khác bán không tốt.")
            format_as_subbullet(median_insight)
        add_spacer()
    
    doc.add_paragraph()  # Spacing
    
    # 8. Official vs third-party
    if stats.get("official_analysis"):
        add_section_title("8. Câu Chuyện: Official Store vs Third-party Sellers")
        official = stats["official_analysis"]
        
        official_story = doc.add_paragraph()
        official_story.add_run("Trên Tiki, có hai loại người bán: Official Store (cửa hàng chính thức) và Third-party Sellers (nhà bán lẻ độc lập). ")
        official_story.add_run("Câu hỏi đặt ra là: Cửa hàng chính thức có thực sự bán tốt hơn không? ")
        
        if official.get("official_products") and official.get("third_party_products"):
            total_products = official["official_products"] + official["third_party_products"]
            official_pct = (official["official_products"] / total_products) * 100
            third_party_pct = (official["third_party_products"] / total_products) * 100
            
            official_story.add_run(f"Dataset cho thấy: ")
            official_story.add_run(f"{official_pct:.1f}%").bold = True
            official_story.add_run(f" sản phẩm từ Official Store, ")
            official_story.add_run(f"{third_party_pct:.1f}%").bold = True
            official_story.add_run(f" từ Third-party.")
        
        format_as_subbullet(official_story)
        
        # So sánh hiệu suất
        if official.get("avg_sales_official") and official.get("avg_sales_third_party"):
            comparison = doc.add_paragraph()
            official_sales = official.get("avg_sales_official") or 0
            third_party_sales = official.get("avg_sales_third_party") or 0
            official_rating = official.get("avg_rating_official") or 0
            third_party_rating = official.get("avg_rating_third_party") or 0
            
            if official_sales > third_party_sales:
                diff = ((official_sales - third_party_sales) / third_party_sales) * 100 if third_party_sales > 0 else 0
                comparison.add_run(f"Thú vị là: Official Store bán tốt hơn trung bình ")
                comparison.add_run(f"{diff:.1f}%").bold = True
                comparison.add_run(f" so với Third-party. ")
            else:
                diff = ((third_party_sales - official_sales) / official_sales) * 100 if official_sales > 0 else 0
                comparison.add_run(f"Điều bất ngờ là: Third-party Sellers bán tốt hơn Official Store trung bình ")
                comparison.add_run(f"{diff:.1f}%").bold = True
                comparison.add_run(f". ")
            
            comparison.add_run("Điều này có thể do Third-party thường có giá cạnh tranh hơn hoặc người tiêu dùng tin tưởng vào các review từ người dùng thực.")
            
            format_as_subbullet(comparison)
        
        # So sánh rating
        if official.get("avg_rating_official") and official.get("avg_rating_third_party"):
            rating_para = doc.add_paragraph()
            official_rating = official.get("avg_rating_official") or 0
            third_party_rating = official.get("avg_rating_third_party") or 0
            
            rating_para.add_run("Về chất lượng (dựa trên rating): ")
            rating_para.add_run(f"Official Store có rating {safe_format(official_rating, '.2f')}/5, ").bold = True
            rating_para.add_run(f"Third-party có {safe_format(third_party_rating, '.2f')}/5. ")
            
            if official_rating > third_party_rating:
                rating_para.add_run("Official Store có ưu thế về chất lượng nhận thức từ người tiêu dùng.")
            else:
                rating_para.add_run("Người tiêu dùng đánh giá Third-party cao hơn, cho thấy sự cạnh tranh lành mạnh.")
            format_as_subbullet(rating_para)
        add_spacer()
    
    doc.add_paragraph()  # Spacing
    
    # 9. Stock
    if stats.get("stock_analysis"):
        add_section_title("9. Câu Chuyện: Tính Sẵn Lòng Bán Hàng")
        stock = stats["stock_analysis"]
        
        stock_story = doc.add_paragraph()
        stock_story.add_run("Một yếu tố quan trọng để người tiêu dùng mua được sản phẩm: hàng phải còn trong kho. ")
        
        if stock.get("in_stock") and stock.get("out_of_stock"):
            total_stock = stock["in_stock"] + stock["out_of_stock"]
            in_stock_pct = (stock["in_stock"] / total_stock) * 100
            out_stock_pct = (stock["out_of_stock"] / total_stock) * 100
            
            stock_story.add_run(f"Dataset cho thấy: ")
            stock_story.add_run(f"{in_stock_pct:.1f}%").bold = True
            stock_story.add_run(f" sản phẩm còn trong kho, ")
            stock_story.add_run(f"{out_stock_pct:.1f}%").bold = True
            stock_story.add_run(f" sản phẩm hết hàng. ")
        
        format_as_subbullet(stock_story)
        
        # So sánh doanh số
        if stock.get("avg_sales_in_stock") and stock.get("avg_sales_out_of_stock"):
            stock_impact = doc.add_paragraph()
            in_stock_sales = stock.get("avg_sales_in_stock") or 0
            out_stock_sales = stock.get("avg_sales_out_of_stock") or 0
            
            stock_impact.add_run("Sản phẩm có trong kho bán được trung bình ")
            stock_impact.add_run(f"{safe_format(in_stock_sales, '.0f')} cái, ").bold = True
            stock_impact.add_run("trong khi sản phẩm hết hàng chỉ bán được ")
            stock_impact.add_run(f"{safe_format(out_stock_sales, '.0f')} cái. ").bold = True
            
            if in_stock_sales > 0 and out_stock_sales > 0:
                diff_ratio = in_stock_sales / out_stock_sales
                stock_impact.add_run(f"Sản phẩm có trong kho bán gấp ")
                stock_impact.add_run(f"{diff_ratio:.1f}x").bold = True
                stock_impact.add_run(" so với hết hàng! Điều này cho thấy sự sẵn sàng của người bán rất quan trọng đến doanh số.")
            format_as_subbullet(stock_impact)
        add_spacer()
    
    doc.add_paragraph()  # Spacing
    
    # 10. Kết luận (retain numbering style from user request but bullet formatting)
    add_section_title("10. Kết Luận: Những Câu Chuyện Đã Kể")
    
    # Tổng hợp câu chuyện
    conclusion = doc.add_paragraph()
    conclusion.add_run("Qua quá trình phân tích dữ liệu, đã khám phá được nhiều insights thú vị về thị trường thương mại điện tử Việt Nam. ")
    conclusion.add_run(f"Với ")
    conclusion.add_run(f"{overview['total_products']:,}").bold = True
    conclusion.add_run(" sản phẩm từ ")
    conclusion.add_run(f"{overview['total_categories']:,}").bold = True
    conclusion.add_run(" danh mục, dataset này phản ánh một thị trường sôi động, đa dạng và đầy tiềm năng.")
    
    for run in conclusion.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # Những phát hiện chính
    doc.add_heading("Những Phát Hiện Chính", 2)
    learnings_intro = doc.add_paragraph()
    learnings_intro.add_run("Từ phân tích dữ liệu, có thể rút ra những phát hiện sau:")
    
    for run in learnings_intro.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(12)
    
    # Tính toán learnings từ dữ liệu
    learnings = []
    
    if prod_stats.get("distinct_brands"):
        learnings.append(f"Thị trường có {prod_stats['distinct_brands']:,} thương hiệu khác nhau, cho thấy sự đa dạng và cạnh tranh trên thị trường.")
    
    if prod_stats.get("discounted_products") and overview.get("total_products"):
        discount_ratio = (prod_stats["discounted_products"] / overview["total_products"]) * 100
        learnings.append(f"Khuyến mãi là một chiến lược phổ biến - có {discount_ratio:.1f}% sản phẩm đang giảm giá.")
    
    if prod_stats.get("avg_rating"):
        learnings.append(f"Người tiêu dùng khá hài lòng với sản phẩm trên thị trường, với điểm đánh giá trung bình {safe_format(prod_stats.get('avg_rating'), '.2f')}/5.0.")
    
    if stats.get("sales_distribution") and stats["sales_distribution"].get("bestsellers"):
        bestseller_count = stats["sales_distribution"]["bestsellers"]
        learnings.append(f"Hiệu ứng '80/20' xuất hiện rõ: chỉ {bestseller_count} sản phẩm ({(bestseller_count/overview.get('total_products', 1)*100):.1f}%) là bestsellers, nhưng họ chiếm lượng bán đáng kể.")
    
    if stats.get("official_analysis"):
        official = stats["official_analysis"]
        if official.get("official_products"):
            official_pct = (official["official_products"] / (official["official_products"] + official.get("third_party_products", 1))) * 100
            learnings.append(f"Sự cân bằng giữa Official Store ({official_pct:.1f}%) và Third-party Sellers tạo ra một thị trường đa dạng và cạnh tranh.")
    
    if stats.get("stock_analysis"):
        stock = stats["stock_analysis"]
        if stock.get("in_stock"):
            in_stock_pct = (stock["in_stock"] / (stock["in_stock"] + stock.get("out_of_stock", 1))) * 100
            learnings.append(f"Tính sẵn lòng: {in_stock_pct:.1f}% sản phẩm còn trong kho. Sản phẩm có trong kho bán rất tốt hơn hết hàng.")
    
    learnings.append(f"Phân hóa giữa bestsellers và sản phẩm thường - không phải tất cả sản phẩm đều bán chạy, thị trường có phân tầng rõ ràng.")
    learnings.append("Người tiêu dùng online không chỉ quan tâm giá, mà còn chất lượng (rating), thương hiệu, và tính sẵn sàng của người bán.")
    
    for learning in learnings:
        p = doc.add_paragraph(learning, style="List Bullet 2")
        for run in p.runs:
            set_vietnamese_font(run)
            run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # Ứng dụng
    doc.add_heading("Ứng Dụng Của Dataset", 2)
    application = doc.add_paragraph()
    application.add_run("Dataset này có thể được sử dụng cho nhiều mục đích nghiên cứu và phân tích:")
    
    for run in application.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(12)
    
    applications = [
        "Nghiên cứu về hành vi mua sắm của người Việt Nam",
        "Phân tích xu hướng thị trường và dự đoán tương lai",
        "So sánh các thương hiệu và sản phẩm",
        "Học tập về Data Engineering và Data Analytics với dữ liệu thực tế",
        "Phát triển các dự án phân tích dữ liệu khác"
    ]
    
    for app in applications:
        p = doc.add_paragraph(app, style="List Bullet 2")
        for run in p.runs:
            set_vietnamese_font(run)
            run.font.size = Pt(12)
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph("---")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = doc.add_paragraph("Tiki Data Pipeline - Data Story Document")
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_text.runs:
        set_vietnamese_font(run)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Lưu file (xóa file cũ nếu tồn tại và đang bị lock)
    try:
        if output_path.exists():
            try:
                output_path.unlink()
            except PermissionError:
                print(f"⚠️  File đang được mở, vui lòng đóng file {output_path} và chạy lại script")
                raise
        doc.save(str(output_path))
        print(f"✅ Đã tạo file docx: {output_path}")
    except PermissionError as e:
        print(f"❌ Không thể ghi file: {e}")
        print(f"💡 Vui lòng đóng file {output_path} nếu đang mở và chạy lại script")
        raise


def main():
    """Hàm chính"""
    # Set UTF-8 encoding for Windows
    if sys.platform == "win32":
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8")
    
    # Parse arguments
    parser = argparse.ArgumentParser(
        description='Xây dựng câu chuyện dữ liệu từ database và tạo file DOCX'
    )
    parser.add_argument(
        '--upload',
        action='store_true',
        help='Tự động upload file lên Google Drive sau khi build'
    )
    parser.add_argument(
        '--folder-id',
        type=str,
        help='ID của folder trên Google Drive để upload file vào',
        default=None
    )
    args = parser.parse_args()
    
    print("=" * 70)
    print("📖 XÂY DỰNG CÂU CHUYỆN DỮ LIỆU - TIKI DATA PIPELINE")
    print("=" * 70)
    
    # Load config
    config = load_env_config()
    print(f"\n📋 Thông tin kết nối:")
    print(f"   Host: {config['host']}")
    print(f"   Port: {config['port']}")
    print(f"   Database: {config['database']}")
    print(f"   User: {config['user']}")
    
    # Kết nối database
    conn = connect_database(config)
    
    try:
        # Phân tích dữ liệu
        print("\n🔍 Bắt đầu phân tích dữ liệu...")
        stats = analyze_database(conn)
        
        # Tạo file docx (luôn ghi đè vào cùng một file)
        output_dir = PROJECT_ROOT / "docs"
        output_dir.mkdir(exist_ok=True)
        output_path = output_dir / "data_story.docx"
        
        print("\n📝 Đang tạo file docx...")
        create_document(stats, output_path)
        
        print("\n" + "=" * 70)
        print("✅ HOÀN TẤT!")
        print("=" * 70)
        print(f"\n📄 File đã được tạo tại: {output_path}")
        if output_path.exists():
            print(f"   Kích thước: {output_path.stat().st_size / 1024:.2f} KB")
        
        # Upload lên Google Drive nếu được yêu cầu
        if args.upload:
            print("\n" + "=" * 70)
            print("📤 UPLOAD LÊN GOOGLE DRIVE")
            print("=" * 70)
            upload_to_google_drive(output_path, args.folder_id)
        
    except Exception as e:
        print(f"\n❌ Lỗi: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
    finally:
        conn.close()
        print("\n🔌 Đã đóng kết nối database")


if __name__ == "__main__":
    main()

